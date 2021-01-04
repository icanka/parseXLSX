from openpyxl import load_workbook
from docx import Document

wb = load_workbook(filename='Financial Sample.xlsx')
dest_filename = 'Financial Sample.xlsx'
ws = wb['Sheet1']
print(type(ws).__name__)


def get_col_names(sheet):
    col_names = {}
    for row in sheet.iter_rows():
        for cell in row:
            if cell.row == 2:
                return col_names
            else:
                col_names[cell.col_idx] = cell.value


col_names = get_col_names(ws)

key_list = list(col_names.keys())

col1 = [ws.cell(i, key_list[0]).row for i in range(2, ws.max_row + 1) if ws.cell(i, key_list[0]).value != 'Government']
print(len(col1))
col2 = [ws.cell(i, key_list[1]).value for i in col1]
col3 = [ws.cell(i, key_list[0]).value for i in col1]
print(len(col2))
zipped = zip(col1, col2)
print(zipped)

document = Document()
document.add_heading('TABLO', level=1)

table = document.add_table(rows=len(col1), cols=2)

# row_count = len(table.rows)
# col_count = len(table.columns)
for i in range(0, len(col1)):
    table.cell(i, 0).text = col3[i]
    table.cell(i, 1).text = col2[i]


# for row in table.rows:
#     for cell in row.cells:
#         cell.text = "hello"

document.save('test.docx')
